"""
WER/CER Evaluation Module
Based on SENSE App Jupyter Notebook
"""
from __future__ import annotations

import re
import json
import docx
import zipfile
from pathlib import Path
from typing import Dict, Optional
from jiwer import wer, cer, wil, process_words
from num2words import num2words


def get_clean_text_from_docx(docx_file: str | Path) -> str:
    """Baca dan bersihkan teks dari file Ground Truth (DOCX/TXT/JSON).

    Nama fungsi dipertahankan untuk kompatibilitas, tapi sekarang mendukung:
    - .docx: dibaca via python-docx
    - .txt: dibaca sebagai plain text
    - .json: mencoba ambil field 'text' atau stringify
    """
    if not docx_file:
        return ""
    
    docx_path = Path(docx_file)

    if not docx_path.exists():
        print(f"  ⚠ SKIP: File Ground Truth tidak ditemukan: {docx_path.name}")
        print(f"  ⚠ Path: {docx_path}")
        return ""

    ext = docx_path.suffix.lower()

    # Plain text ground truth
    if ext == ".txt":
        try:
            return docx_path.read_text(encoding="utf-8", errors="ignore")
        except Exception as e:
            print(f"  ⚠ SKIP: Gagal membaca Ground Truth TXT: {docx_path.name}")
            print(f"  ⚠ Reason: {e}")
            return ""

    # JSON ground truth (optional)
    if ext == ".json":
        try:
            obj = json.loads(docx_path.read_text(encoding="utf-8", errors="ignore"))
            if isinstance(obj, dict) and "text" in obj:
                return str(obj["text"])
            return json.dumps(obj, ensure_ascii=False)
        except Exception as e:
            print(f"  ⚠ SKIP: Gagal membaca Ground Truth JSON: {docx_path.name}")
            print(f"  ⚠ Reason: {e}")
            return ""
    
    try:
        # DOCX ground truth
        doc = docx.Document(str(docx_path))
    except (zipfile.BadZipFile, Exception) as e:
        # JIKA ERROR: Cetak pesan warning, tapi JANGAN stop program
        # Catatan: python-docx juga melempar error serupa jika file bukan DOCX valid.
        print(f"  ⚠ SKIP: File Ground Truth tidak bisa dibaca sebagai DOCX: {docx_path.name}")
        print(f"  ⚠ Reason: {e}")
        return ""  # Kembalikan teks kosong
    
    full_text = []
    
    # Prioritas Tabel
    has_table_content = False
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join([cell.text for cell in row.cells])
            if len(row_text.strip()) > 5:
                full_text.append(row_text)
                has_table_content = True
    
    if not has_table_content:
        for para in doc.paragraphs:
            full_text.append(para.text)
    
    raw_text = "\n".join(full_text)
    
    # Cleaning patterns
    garbage_patterns = [
        r"Interviewer\s*[:|]", r"Narasumber\s*[:|]", r"Responden\s*[:|]",
        r"Wawancara\s*[:|]", r"Transkripter\s*[:|]", r"Identitas.*",
        r"Nama\s*[:|]", r"Umur\s*[:|]", r"Jenis Kelamin\s*[:|]",
        r"Tanggal.*", r"Tempat.*", r"Bagian Teks", r"Pembicara",
        r"^=\s*", r"\[.*?\]"
    ]
    
    for pat in garbage_patterns:
        raw_text = re.sub(pat, " ", raw_text, flags=re.IGNORECASE)
    
    return raw_text


def normalize_aggressive(text: str) -> str:
    """Normalisasi teks untuk evaluasi WER/CER"""
    
    # Hapus karakter non-alphanumeric
    text = re.sub(r"[^a-zA-Z0-9\s]", " ", text)
    
    # Convert numbers to words (Indonesian)
    try:
        text = re.sub(r'\d+', lambda x: num2words(int(x.group()), lang='id'), text)
    except:
        pass
    
    text = text.lower()
    
    # Hapus filler words
    fillers = [
        "eee", "hmm", "anu", "nya", "sih", "dong", "tuh", "mah", 
        "yak", "yah", "eung", "oh", "iya", "ya", "nah", "oke"
    ]
    
    words = text.split()
    clean_words = []
    for w in words:
        # Hapus karakter berulang (e.g., "yaaa" -> "ya")
        w = re.sub(r'(.)\1{2,}', r'\1', w)
        if w not in fillers and len(w) > 1:
            clean_words.append(w)
    
    return " ".join(clean_words)


def calculate_wer_cer(reference: str, hypothesis: str) -> Dict[str, float]:
    """
    Calculate WER, CER, and WIL metrics
    
    Args:
        reference: Ground truth text
        hypothesis: Predicted/transcribed text
    
    Returns:
        Dictionary with metrics
    """
    # Normalize both texts
    ref = normalize_aggressive(reference)
    hyp = normalize_aggressive(hypothesis)
    
    if not ref or not hyp:
        return {
            "wer": 0.0,
            "cer": 0.0,
            "wil": 0.0,
            "word_count_ref": len(ref.split()) if ref else 0,
            "word_count_hyp": len(hyp.split()) if hyp else 0,
            "char_count_ref": len(ref) if ref else 0,
            "char_count_hyp": len(hyp) if hyp else 0
        }
    
    try:
        wer_score = wer(ref, hyp)
        cer_score = cer(ref, hyp)
        wil_score = wil(ref, hyp)
    except Exception as e:
        print(f"Error calculating metrics: {e}")
        return {
            "wer": 0.0,
            "cer": 0.0,
            "wil": 0.0,
            "word_count_ref": len(ref.split()),
            "word_count_hyp": len(hyp.split()),
            "char_count_ref": len(ref),
            "char_count_hyp": len(hyp),
            "error": str(e)
        }
    
    return {
        "wer": wer_score,
        "cer": cer_score,
        "wil": wil_score,
        "word_count_ref": len(ref.split()),
        "word_count_hyp": len(hyp.split()),
        "char_count_ref": len(ref),
        "char_count_hyp": len(hyp)
    }


def evaluate_transcript(ground_truth_path: str | Path, hypothesis_text: str) -> Optional[Dict]:
    """
    Evaluate transcript against ground truth
    
    Args:
        ground_truth_path: Path to ground truth DOCX file
        hypothesis_text: Transcribed text to evaluate
    
    Returns:
        Evaluation metrics or None if ground truth invalid
    """
    gt_text = get_clean_text_from_docx(ground_truth_path)
    
    if not gt_text or len(gt_text.strip()) < 10:
        return None
    
    metrics = calculate_wer_cer(gt_text, hypothesis_text)
    
    return {
        "ground_truth_file": Path(ground_truth_path).name,
        **metrics
    }


def match_audio_with_gt(audio_path: Path, gt_files: Dict[str, Path]) -> Optional[Path]:
    """
    Match audio dengan ground truth berdasarkan nama
    
    Args:
        audio_path: Path to audio file
        gt_files: Dictionary of {key: path} for ground truth files
    
    Returns:
        Path to matching ground truth file or None
    """
    audio_name = audio_path.stem
    audio_clean = audio_name.lower().replace('_', '').replace(' ', '').replace('-', '')
    
    # Strategy 1: Exact match
    if audio_name in gt_files:
        return gt_files[audio_name]
    
    # Strategy 2: Case-insensitive
    for gt_key, gt_path in gt_files.items():
        if gt_key.lower() == audio_name.lower():
            return gt_path
    
    # Strategy 3: Contains match
    for gt_key, gt_path in gt_files.items():
        if gt_key.lower() in audio_name.lower() or audio_name.lower() in gt_key.lower():
            return gt_path
    
    # Strategy 4: Cleaned match
    for gt_key, gt_path in gt_files.items():
        gt_clean = gt_key.lower().replace('_', '').replace(' ', '').replace('-', '')
        if gt_clean == audio_clean or gt_clean in audio_clean or audio_clean in gt_clean:
            return gt_path
    
    return None


